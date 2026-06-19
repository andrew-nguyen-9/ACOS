from __future__ import annotations

import io

import pytest
from docx import Document

from backend.services.cover_letter.docx_exporter import CoverLetterDOCXExporter


@pytest.fixture
def exporter() -> CoverLetterDOCXExporter:
    return CoverLetterDOCXExporter()


def test_export_returns_bytes(exporter):
    result = exporter.export("Dear Hiring Manager,\n\nI am excited.", "Engineer", "Acme")
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_export_produces_valid_docx(exporter):
    result = exporter.export("Dear Hiring Manager,\n\nI am excited.", "Engineer", "Acme")
    doc = Document(io.BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    assert any("Dear Hiring Manager" in t for t in texts)


def test_export_multiline_produces_paragraphs(exporter):
    text = "Line one\nLine two\nLine three"
    result = exporter.export(text, "Developer", "Corp")
    doc = Document(io.BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    assert any("Line one" in t for t in texts)
    assert any("Line two" in t for t in texts)
    assert any("Line three" in t for t in texts)


def test_export_empty_text_returns_bytes(exporter):
    result = exporter.export("", "Dev", "Corp")
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_export_never_raises_on_unusual_input(exporter):
    """export() contract: never raises, always returns bytes."""
    result = exporter.export("Hello\x00World", "Title", "Company")
    assert isinstance(result, bytes)
