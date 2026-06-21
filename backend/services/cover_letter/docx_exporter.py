from __future__ import annotations

import io
import logging
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)

_FONT = "Arial"
_SZ_HEADER = Pt(14)
_SZ_CONTACT = Pt(11)
_SZ_BODY = Pt(11)


class CoverLetterDOCXExporter:
    """Exports cover letter text to a formatted business-letter DOCX stream.

    Produces: candidate header block, date, recipient/subject block, body, closing.
    Never raises — falls back to plain text DOCX on any error.
    """

    def export(
        self,
        text: str,
        job_title: str,
        company: str,
        contact_info: dict | None = None,
    ) -> bytes:
        """Render text into a formatted business-letter DOCX and return raw bytes.

        Args:
            text: Plain-text cover letter body (newlines delimit paragraphs).
            job_title: Role title — used in the Re: subject line.
            company: Company name — used in the recipient block.
            contact_info: Dict with name, location, phone, email, linkedin keys.

        Returns:
            Raw DOCX bytes; never raises.
        """
        try:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Pt(54)
                section.bottom_margin = Pt(54)
                section.left_margin = Pt(72)
                section.right_margin = Pt(72)

            info = contact_info or {}
            self._add_candidate_header(doc, info)
            self._add_date(doc)
            self._add_recipient_block(doc, company, job_title)
            self._add_body(doc, text)
            self._add_closing(doc, info.get("name", ""))

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as exc:
            logger.error(
                "cl_docx_exporter: export failed for '%s' @ '%s': %s",
                job_title,
                company,
                exc,
            )
            return self._fallback(text)

    # ── Candidate header ──────────────────────────────────────────────────────

    def _add_candidate_header(self, doc: object, info: dict) -> None:
        name = info.get("name", "")
        if name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(name.upper())
            r.font.name = _FONT
            r.font.size = _SZ_HEADER
            r.bold = True

        parts = [
            info.get("location", ""),
            info.get("phone", ""),
            info.get("email", ""),
            info.get("linkedin", ""),
        ]
        contact_line = " | ".join(x for x in parts if x)
        if contact_line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(contact_line)
            r.font.name = _FONT
            r.font.size = _SZ_CONTACT

        doc.add_paragraph()

    def _add_date(self, doc: object) -> None:
        p = doc.add_paragraph()
        r = p.add_run(date.today().strftime("%B %d, %Y"))
        r.font.name = _FONT
        r.font.size = _SZ_BODY
        doc.add_paragraph()

    def _add_recipient_block(self, doc: object, company: str, job_title: str) -> None:
        for line in filter(None, ["Hiring Manager", company]):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = _FONT
            r.font.size = _SZ_BODY
        if job_title:
            p = doc.add_paragraph()
            r = p.add_run(f"Re: {job_title}")
            r.font.name = _FONT
            r.font.size = _SZ_BODY
            r.bold = True
        doc.add_paragraph()

    # ── Body ─────────────────────────────────────────────────────────────────

    def _add_body(self, doc: object, text: str) -> None:
        for line in text.split("\n"):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = _FONT
            r.font.size = _SZ_BODY

    # ── Closing ───────────────────────────────────────────────────────────────

    def _add_closing(self, doc: object, name: str) -> None:
        doc.add_paragraph()
        for line in ["Sincerely,", "", name]:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = _FONT
            r.font.size = _SZ_BODY

    # ── Fallback ──────────────────────────────────────────────────────────────

    def _fallback(self, text: str) -> bytes:
        try:
            doc = Document()
            doc.add_paragraph(text)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as inner_exc:
            logger.error("cl_docx_exporter: fallback also failed: %s", inner_exc)
            return b""
