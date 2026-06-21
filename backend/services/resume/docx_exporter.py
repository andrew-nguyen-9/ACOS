from __future__ import annotations

import io
import logging
from pathlib import Path

from docx import Document  # type: ignore[name-defined]
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

_WEAK_MARKER = "⚠ [REVIEW REQUIRED]"

_TEMPLATE_DIR = Path(__file__).parents[4] / "mock-designs" / "resume-templates"

_TEMPLATE_FILES: dict[str, str] = {
    "software": "Yale-College-General-Template-v.1.docx",
    "ai": "Yale-College-General-Template-v.1.docx",
    "data_analytics": "Resume-Fewer-than-10-yrs.docx",
    "product": "Resume-Fewer-than-10-yrs.docx",
    "consulting": "tMPA_Application_Resume_Template.docx",
    "healthcare": "tMPA_Application_Resume_Template.docx",
}

_FONT = "Arial"
_SZ_NAME = Pt(14)
_SZ_CONTACT = Pt(11)
_SZ_BODY = Pt(10)
_ORANGE = RGBColor(0xFF, 0x80, 0x00)


def _add_bottom_border(para: object) -> None:
    pPr = para.paragraph_format._element.get_or_add_pPr()  # type: ignore[attr-defined]
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


class ResumeDOCXExporter:
    def export(
        self,
        content_json: dict,
        template_name: str,
        contact_info: dict | None = None,
    ) -> bytes:
        """Export resume to DOCX using the visual template matching template_name.

        Args:
            content_json: Resume dict (header, experiences, skills, projects, education).
            template_name: Registered template name — determines which .docx style is used.
            contact_info: Override for name/contact. Falls back to content_json["header"].

        Returns:
            DOCX bytes; never raises.
        """
        try:
            doc = self._load_template(template_name)
            self._clear_body(doc)

            header = contact_info or content_json.get("header", {})
            self._add_header(doc, header)

            experiences = content_json.get("experiences", [])
            skills = content_json.get("skills", [])
            projects = content_json.get("projects", [])
            education = content_json.get("education", [])

            # Yale template puts Education first (academic → career order)
            if _TEMPLATE_FILES.get(template_name, "").startswith("Yale"):
                self._add_education(doc, education)
                self._add_experiences(doc, experiences)
                self._add_projects(doc, projects)
                self._add_skills(doc, skills)
            else:
                self._add_experiences(doc, experiences)
                self._add_skills(doc, skills)
                self._add_projects(doc, projects)
                self._add_education(doc, education)

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as exc:
            logger.error("resume_docx_exporter: export failed: %s", exc)
            return self._fallback(content_json)

    # ── Template loading ──────────────────────────────────────────────────────

    def _load_template(self, template_name: str) -> object:
        fname = _TEMPLATE_FILES.get(template_name)
        if fname:
            path = _TEMPLATE_DIR / fname
            if path.exists():
                return Document(str(path))  # type: ignore[operator]
        return Document()  # type: ignore[operator]

    def _clear_body(self, doc: object) -> None:
        for para in list(doc.paragraphs):  # type: ignore[attr-defined]
            para._element.getparent().remove(para._element)

    # ── Header ────────────────────────────────────────────────────────────────

    def _add_header(self, doc: object, header: dict) -> None:
        name = header.get("name", "")
        if name:
            p = doc.add_paragraph()  # type: ignore[attr-defined]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(name.upper())  # type: ignore[attr-defined]
            r.font.name = _FONT
            r.font.size = _SZ_NAME
            r.bold = True

        parts = [
            header.get("location", ""),
            header.get("phone", ""),
            header.get("email", ""),
            header.get("linkedin", ""),
            header.get("github", ""),
        ]
        contact_line = " | ".join(x for x in parts if x)
        if contact_line:
            p = doc.add_paragraph()  # type: ignore[attr-defined]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(contact_line)  # type: ignore[attr-defined]
            r.font.name = _FONT
            r.font.size = _SZ_CONTACT

        doc.add_paragraph()  # type: ignore[attr-defined]

    # ── Section helpers ───────────────────────────────────────────────────────

    def _section_header(self, doc: object, title: str) -> None:
        p = doc.add_paragraph()  # type: ignore[attr-defined]
        r = p.add_run(title.upper())  # type: ignore[attr-defined]
        r.font.name = _FONT
        r.font.size = _SZ_BODY
        r.bold = True
        _add_bottom_border(p)

    def _run(self, para: object, text: str, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
        r = para.add_run(text)  # type: ignore[attr-defined]
        r.font.name = _FONT
        r.font.size = _SZ_BODY
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color

    # ── Sections ──────────────────────────────────────────────────────────────

    def _add_experiences(self, doc: object, experiences: list[dict]) -> None:
        if not experiences:
            return
        self._section_header(doc, "Experience")
        for exp in experiences:
            p = doc.add_paragraph()  # type: ignore[attr-defined]
            self._run(p, exp.get("company", ""), bold=True)
            dates = exp.get("dates", "")
            if dates:
                self._run(p, f"\t{dates}")
            if exp.get("title"):
                p2 = doc.add_paragraph()  # type: ignore[attr-defined]
                self._run(p2, exp["title"], italic=True)
            for bullet in exp.get("bullets", []):
                self._add_bullet(doc, bullet)
            doc.add_paragraph()  # type: ignore[attr-defined]

    def _add_bullet(self, doc: object, bullet: object) -> None:
        if isinstance(bullet, dict):
            text = bullet.get("text", "")
            confidence = bullet.get("confidence", "verified")
        else:
            text = str(bullet)
            confidence = "verified"
        try:
            p = doc.add_paragraph(style="List Bullet")  # type: ignore[attr-defined]
        except Exception:
            p = doc.add_paragraph()  # type: ignore[attr-defined]
        is_weak = confidence == "weak_inference"
        self._run(
            p,
            f"{_WEAK_MARKER} {text}" if is_weak else text,
            color=_ORANGE if is_weak else None,
        )

    def _add_skills(self, doc: object, skills: list[str]) -> None:
        if not skills:
            return
        self._section_header(doc, "Skills")
        p = doc.add_paragraph()  # type: ignore[attr-defined]
        self._run(p, " · ".join(skills))
        doc.add_paragraph()  # type: ignore[attr-defined]

    def _add_projects(self, doc: object, projects: list[dict]) -> None:
        if not projects:
            return
        self._section_header(doc, "Projects")
        for proj in projects:
            p = doc.add_paragraph()  # type: ignore[attr-defined]
            self._run(p, proj.get("name", ""), bold=True)
            if proj.get("description"):
                self._run(p, f" — {proj['description']}")
            if proj.get("tech"):
                self._run(p, f" [{proj['tech']}]")
        doc.add_paragraph()  # type: ignore[attr-defined]

    def _add_education(self, doc: object, education: list[dict]) -> None:
        if not education:
            return
        self._section_header(doc, "Education")
        for edu in education:
            p = doc.add_paragraph()  # type: ignore[attr-defined]
            self._run(p, edu.get("institution", ""), bold=True)
            if edu.get("dates"):
                self._run(p, f"\t{edu['dates']}")
            if edu.get("degree"):
                p2 = doc.add_paragraph()  # type: ignore[attr-defined]
                self._run(p2, edu["degree"])
        doc.add_paragraph()  # type: ignore[attr-defined]

    # ── Fallback ──────────────────────────────────────────────────────────────

    def _fallback(self, content_json: dict) -> bytes:
        try:
            doc = Document()  # type: ignore[operator]
            doc.add_paragraph("Resume export error — please retry.")
            for exp in content_json.get("experiences", []):
                doc.add_paragraph(f"{exp.get('title', '')} at {exp.get('company', '')}")
                for b in exp.get("bullets", []):
                    t = b.get("text", "") if isinstance(b, dict) else str(b)
                    doc.add_paragraph(f"• {t}")
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:
            return b""
