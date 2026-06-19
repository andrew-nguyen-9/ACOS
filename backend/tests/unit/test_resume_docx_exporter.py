import io
import pytest
from docx import Document
from backend.services.resume.docx_exporter import ResumeDOCXExporter

_SAMPLE_CONTENT = {
    "summary": "Senior Data Engineer with 5 years of Python expertise.",
    "experiences": [
        {
            "title": "Data Engineer",
            "company": "Acme Corp",
            "dates": "2022–2024",
            "bullets": [
                {"text": "Built ETL pipeline in Python.", "evidence_id": "b1", "confidence": "verified"},
                {"text": "Maybe led a team.", "evidence_id": "b2", "confidence": "weak_inference"},
            ],
        }
    ],
    "skills": ["Python", "SQL", "ETL"],
    "projects": [{"name": "ACOS", "description": "Career OS", "tech": "Python/FastAPI", "evidence_id": "p1"}],
    "education": [],
}


def test_export_returns_bytes():
    exporter = ResumeDOCXExporter()
    result = exporter.export(_SAMPLE_CONTENT, "software")
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_export_produces_valid_docx():
    exporter = ResumeDOCXExporter()
    result = exporter.export(_SAMPLE_CONTENT, "software")
    doc = Document(io.BytesIO(result))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Data Engineer" in text
    assert "Acme Corp" in text


def test_export_flags_weak_inference():
    exporter = ResumeDOCXExporter()
    result = exporter.export(_SAMPLE_CONTENT, "software")
    doc = Document(io.BytesIO(result))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "⚠" in text or "[REVIEW REQUIRED]" in text


def test_export_includes_skills():
    exporter = ResumeDOCXExporter()
    result = exporter.export(_SAMPLE_CONTENT, "software")
    doc = Document(io.BytesIO(result))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Python" in text
